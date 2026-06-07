from copy import deepcopy
from io import BytesIO
import json

import yaml

from app.services.document_conversion_service import convert_document


INTERNAL_KEYS = {
    "content_block_id",
    "source_evidence_ids",
    "source_paragraph_ids",
    "paragraph_id",
    "paragraph_ids",
    "traceability_index",
}
EXPORT_EXTENSIONS = {
    "yaml": "yaml",
    "markdown": "md",
    "doc": "doc",
    "docx": "docx",
    "pdf": "pdf",
    "txt": "txt",
    "clean_json": "json",
}
EXPORT_CONTENT_TYPES = {
    "yaml": "application/x-yaml; charset=utf-8",
    "markdown": "text/markdown; charset=utf-8",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "txt": "text/plain; charset=utf-8",
    "clean_json": "application/json; charset=utf-8",
}
BLOCK_TYPE_LABELS = {
    "action": "动作",
    "dialogue": "对白",
    "narration": "旁白",
    "transition": "转场",
    "note": "备注",
    "parenthetical": "表演提示",
    "voiceover": "画外音",
    "description": "描述",
    "sound": "声音",
    "character": "人物",
    "shot": "镜头",
}


def _remove_internal(value):
    if isinstance(value, dict):
        return {k: _remove_internal(v) for k, v in value.items() if k not in INTERNAL_KEYS}
    if isinstance(value, list):
        return [_remove_internal(item) for item in value]
    return value


def to_user_clean_json(internal: dict) -> dict:
    return _remove_internal(deepcopy(internal))


def to_yaml_preview(internal_or_clean: dict) -> str:
    clean = to_user_clean_json(internal_or_clean)
    return yaml.safe_dump(clean, allow_unicode=True, sort_keys=False)


def serialize_export(internal: dict, export_format: str) -> str | bytes:
    clean = to_user_clean_json(internal)
    if export_format == "yaml":
        return to_yaml_preview(clean)
    if export_format == "clean_json":
        return json.dumps(clean, ensure_ascii=False, indent=2)
    if export_format in {"markdown", "txt"}:
        return to_yaml_preview(clean)
    if export_format == "docx":
        return _serialize_docx(clean)
    if export_format == "doc":
        return convert_document(_serialize_docx(clean), ".docx", ".doc")
    if export_format == "pdf":
        return convert_document(_serialize_docx(clean), ".docx", ".pdf")
    raise ValueError(f"unsupported export format: {export_format}")


def _serialize_docx(clean: dict) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(str(clean.get("title") or "Script"), level=1)
    for scene in clean.get("scenes") or []:
        document.add_heading(str(scene.get("title") or scene.get("scene_id") or "Scene"), level=2)
        for line in _format_scene_metadata(scene):
            document.add_paragraph(line)
        for block in scene.get("content_blocks") or []:
            line = _format_content_block(block)
            if line:
                document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _format_scene_metadata(scene: dict) -> list[str]:
    lines = [
        f"场景编号：{_string_value(scene.get('scene_id')) or '待定'}",
        f"标题：{_string_value(scene.get('title')) or '待定'}",
    ]
    scene_info = _string_value(scene.get("scene_info"))
    if scene_info:
        lines.append(f"场景信息：{scene_info}")
    characters = _join_values(scene.get("characters"))
    if characters:
        lines.append(f"出场人物：{characters}")
    scene_purpose = _string_value(scene.get("scene_purpose"))
    if scene_purpose:
        lines.append(f"场景目的：{scene_purpose}")
    core_conflict = _string_value(scene.get("core_conflict"))
    if core_conflict:
        lines.append(f"核心冲突：{core_conflict}")
    return lines


def _format_content_block(block: dict) -> str | None:
    text = _string_value(block.get("text"))
    if not text:
        return None

    block_type = _string_value(block.get("type") or block.get("block_type"))
    if block_type == "dialogue":
        speaker = _string_value(block.get("speaker")) or "未指定角色"
        return f"{speaker}：{text}"

    label = _block_type_label(block_type)
    return f"{label}：{text}"


def _block_type_label(block_type: str) -> str:
    return BLOCK_TYPE_LABELS.get(block_type, block_type or "内容")


def _join_values(value) -> str:
    if not isinstance(value, list):
        return ""
    return "、".join(str(item).strip() for item in value if str(item).strip())


def _string_value(value) -> str:
    if value is None:
        return ""
    return str(value).strip()
