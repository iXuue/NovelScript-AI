import base64
from io import BytesIO

import pytest
from docx import Document

from app.services.document_conversion_service import UnsupportedDocumentTypeError
from app.services.input_adapter import normalize_to_markdown

TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_bytes_with_picture() -> bytes:
    document = Document()
    document.add_paragraph("Chapter 1")
    document.add_picture(BytesIO(TINY_PNG))
    document.add_paragraph("The old door opens.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_markdown_upload_strips_image_references_but_keeps_links():
    markdown = normalize_to_markdown(
        "novel.md",
        b"""Chapter 1

![Cover](cover.png)
<img src="map.webp" alt="Map">
![[diagram.jpg]]
![Ref][cover]
[cover]: cover.png "Cover"

Keep [source](https://example.com) link.
""",
    )

    assert "Chapter 1" in markdown
    assert "Keep [source](https://example.com) link." in markdown
    assert "![Cover]" not in markdown
    assert "<img" not in markdown
    assert "![[diagram.jpg]]" not in markdown
    assert "[cover]:" not in markdown


def test_docx_upload_ignores_inline_picture_and_keeps_text():
    markdown = normalize_to_markdown("novel.docx", _docx_bytes_with_picture())

    assert markdown == "Chapter 1\n\nThe old door opens."


def test_doc_upload_is_converted_to_docx_before_markdown_extraction(monkeypatch):
    calls = []
    converted = _docx_bytes("Chapter 1", "The old door opens.")

    def fake_convert_document(content: bytes, source_suffix: str, target_suffix: str) -> bytes:
        calls.append((content, source_suffix, target_suffix))
        return converted

    monkeypatch.setattr("app.services.input_adapter.convert_document", fake_convert_document)

    markdown = normalize_to_markdown("novel.doc", b"legacy-doc-bytes")

    assert markdown == "Chapter 1\n\nThe old door opens."
    assert calls == [(b"legacy-doc-bytes", ".doc", ".docx")]


def test_unsupported_upload_extension_raises_typed_error():
    with pytest.raises(UnsupportedDocumentTypeError):
        normalize_to_markdown("novel.xlsx", b"not text")
